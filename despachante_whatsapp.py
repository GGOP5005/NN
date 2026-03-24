import os
import time
import re
import sys
import json
import unicodedata
import subprocess
from datetime import datetime, timedelta
from colorama import init, Fore, Style
from playwright.sync_api import sync_playwright
from google.oauth2.service_account import Credentials
from googleapiclient.discovery import build
from google import genai
from google.genai import types

# Importação para manipulação de Word e conversão para PDF
try:
    from docx import Document
    from docx2pdf import convert
except ImportError:
    print(Fore.RED + "❌ Bibliotecas ausentes! Execute no terminal: pip install python-docx docx2pdf")
    sys.exit()

# Importação para o motor de áudio gratuito
try:
    import speech_recognition as sr
    from pydub import AudioSegment
except ImportError:
    print(Fore.RED + "❌ Bibliotecas de áudio ausentes! Execute no terminal: pip install SpeechRecognition pydub")
    sys.exit()

from config import BASE_DIR, PLANILHA_ID, LISTA_CHAVES_GEMINI, ROTEAMENTO_PORTOS
from buscador_pdfs import encontrar_pasta_container
from sheets_api import executar_com_resiliencia_infinita, MAPA_MESES, get_col_letter

init(autoreset=True)

CREDS_PATH = os.path.join(BASE_DIR, "credentials.json")
SCOPES = ["https://www.googleapis.com/auth/spreadsheets"]

# ==========================================================
# 🧠 CONFIGURAÇÕES DA IA E CONTATOS AUTORIZADOS
# ==========================================================
CONTATOS_DIRETORIA = [
    "Douglas NN",
    "Financeiro - Norte Nordeste", 
    "Nivaldo - Norte Nordeste",
    "Gabriel - Norte Nordeste",
    "Rivaldo - Norte Nordeste"
]

# CORREÇÃO BUG 7: Cache da aba com controle de mês para invalidação automática
# Antes: ABA_CACHEADA nunca era limpa, então se o robô ligasse em março
# e ficasse rodando até abril, continuava lendo/escrevendo na aba de março
ABA_CACHEADA = None
MES_ABA_CACHEADA = None  # Guarda o mês em que a aba foi cacheada
MODELO_CEREBRO_CACHE = None

def obter_modelo_cerebro(client):
    global MODELO_CEREBRO_CACHE
    if MODELO_CEREBRO_CACHE: return MODELO_CEREBRO_CACHE
    try:
        models = client.models.list()
        for m in models:
            if '2.5-flash' in m.name.lower():
                MODELO_CEREBRO_CACHE = m.name.split('/')[-1]
                return MODELO_CEREBRO_CACHE
    except: pass
    return "gemini-2.5-flash"

# ==========================================================
# 🛡️ ESTERILIZADOR DE TEXTO (ANTI-ERRO DO WORD)
# ==========================================================
def limpar_texto_word(texto):
    """Remove emojis, quebras de linha e caracteres invisíveis que travam o Word."""
    if not texto: return ""
    t = str(texto).replace('\n', ' ').replace('\r', ' ')
    t = ''.join(c for c in t if c <= '\uFFFF') # Remove emojis
    t = ''.join(c for c in t if unicodedata.category(c)[0] != 'C') # Remove controle
    return t.upper().strip()

# ==========================================================
# 🎧 MÓDULO DE TRANSCRIÇÃO (100% LOCAL E GRATUITO)
# ==========================================================
def transcrever_audio_local(caminho_ogg):
    caminho_wav = caminho_ogg.replace(".ogg", ".wav")
    try:
        print(Fore.MAGENTA + "      🎧 Convertendo formato do áudio...")
        audio = AudioSegment.from_file(caminho_ogg, format="ogg")
        audio.export(caminho_wav, format="wav")
        print(Fore.MAGENTA + "      🎧 Transcrevendo áudio (Modo Gratuito)...")
        recognizer = sr.Recognizer()
        with sr.AudioFile(caminho_wav) as source:
            audio_data = recognizer.record(source)
            texto = recognizer.recognize_google(audio_data, language="pt-BR")
        return texto.strip()
    except sr.UnknownValueError:
        print(Fore.YELLOW + "      ⚠️ O áudio está vazio ou não foi possível entender as palavras.")
        return ""
    except Exception as e:
        print(Fore.RED + f"      ❌ Erro na transcrição local: {e}")
        return ""
    finally:
        try:
            if os.path.exists(caminho_ogg): os.remove(caminho_ogg)
            if os.path.exists(caminho_wav): os.remove(caminho_wav)
        except: pass

# ==========================================================
# 📄 TÁTICA HUMANA: COPIAR ARQUIVO (CTRL+C) VIA WINDOWS
# ==========================================================
def copiar_arquivos_windows(arquivos):
    if isinstance(arquivos, str): arquivos = [arquivos]
    caminhos = []
    for arq in arquivos:
        abs_path = os.path.abspath(arq).replace("'", "''")
        caminhos.append(f"'{abs_path}'")
    lista_ps = ",".join(caminhos)
    comando = f'powershell -command "Set-Clipboard -Path {lista_ps}"'
    subprocess.run(comando, shell=True, creationflags=0x08000000 if sys.platform == 'win32' else 0)

# ==========================================================
# 📄 GERADOR DE PROPOSTAS COMERCIAIS
# ==========================================================
def gerar_documento_proposta(cliente, cnpj, origem, destino, valor, att, advaloren, triagem, pedagio, prazo):
    caminho_template = os.path.join(BASE_DIR, "TEMPLATE_PROPOSTA.docx")
    if not os.path.exists(caminho_template):
        return None, "O arquivo 'TEMPLATE_PROPOSTA.docx' não foi encontrado."

    try:
        doc = Document(caminho_template)
        data_atual = datetime.now().strftime("%d/%m/%Y")
        data_validade = (datetime.now() + timedelta(days=30)).strftime("%d/%m/%Y")
        
        # Limpeza severa de variáveis
        cliente_c = limpar_texto_word(cliente)
        cnpj_c = limpar_texto_word(cnpj)
        
        cliente_assinatura = cliente_c
        palavras = cliente_c.split()
        if len(palavras) >= 2: cliente_assinatura = f"{palavras[0]} {palavras[1]}"
        elif len(palavras) == 1: cliente_assinatura = palavras[0]

        substituicoes = {
            "[CLIENTE_CNPJ]": f"{cliente_c} - CNPJ: {cnpj_c}" if cnpj_c and cnpj_c != "A DEFINIR" else cliente_c,
            "[CLIENTE]": cliente_c, 
            "[ASSINATURA]": cliente_assinatura,
            "[CNPJ]": cnpj_c,
            "[ORIGEM]": limpar_texto_word(origem),
            "[DESTINO]": limpar_texto_word(destino), 
            "[VALOR]": limpar_texto_word(valor),
            "[ATT]": limpar_texto_word(att), 
            "[ADVALOREN]": limpar_texto_word(advaloren),
            "[TRIAGEM]": limpar_texto_word(triagem), 
            "[PEDAGIO]": limpar_texto_word(pedagio),
            "[PRAZO]": limpar_texto_word(prazo),
            "[DATA]": data_atual, 
            "[VALIDADE]": data_validade
        }

        for p in doc.paragraphs:
            for key, val in substituicoes.items():
                if key in p.text:
                    for run in p.runs:
                        if key in run.text: run.text = run.text.replace(key, val)
                    if key in p.text: p.text = p.text.replace(key, val)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key, val in substituicoes.items():
                            if key in p.text:
                                for run in p.runs:
                                    if key in run.text: run.text = run.text.replace(key, val)
                                if key in p.text: p.text = p.text.replace(key, val)

        pasta_saida = os.path.join(BASE_DIR, "Propostas_Geradas")
        os.makedirs(pasta_saida, exist_ok=True)
        nome_arquivo_docx = f"Proposta_NN_{re.sub(r'[^A-Za-z0-9]', '', cliente_c)}_{int(time.time())}.docx"
        nome_arquivo_pdf = f"Proposta_NN_{re.sub(r'[^A-Za-z0-9]', '', cliente_c)}_{int(time.time())}.pdf"
        caminho_docx = os.path.join(pasta_saida, nome_arquivo_docx)
        caminho_pdf = os.path.join(pasta_saida, nome_arquivo_pdf)
        
        doc.save(caminho_docx)
        time.sleep(1) 
        
        print(Fore.CYAN + "      🔄 Convertendo Proposta para PDF...")
        try:
            convert(caminho_docx, caminho_pdf)
            time.sleep(2) 
            return caminho_pdf, None
        except Exception as e_pdf:
            print(Fore.YELLOW + "      ⚠️ Falha no Word. Limpando memória e tentando segunda vez...")
            os.system("taskkill /F /IM WINWORD.EXE >nul 2>&1")
            time.sleep(3)
            try:
                convert(caminho_docx, caminho_pdf)
                time.sleep(2) 
                return caminho_pdf, None
            except Exception as e2:
                return None, f"Erro ao converter PDF: {e2}"
            
    except Exception as e:
        return None, f"Erro interno Word: {e}"

# ==========================================================
# 📄 GERADOR MESTRE DE ORDEM DE COLETA
# ==========================================================
def gerar_ordem_coleta(empresa, container, di, lote, bl, motorista, cpf, cavalo, carreta):
    empresa_upper = limpar_texto_word(empresa)
    data_atual = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
    
    if "VAL" in empresa_upper or "MOTOS" in empresa_upper:
        nome_arq_base = "VAL_MOTOS"
        # CORREÇÃO BUG 1: Nome correto do template com underscore entre VAL e MOTOS
        # Antes: "TEMPLATE_COLETA_VALMOTOS.docx" (sem underscore) — arquivo não existia
        caminho_template = os.path.join(BASE_DIR, "TEMPLATE_COLETA_VAL_MOTOS.docx")
        substituicoes = {
            "[DI]": limpar_texto_word(di), "[LOTE]": limpar_texto_word(lote), "[BL]": limpar_texto_word(bl),
            "[MOTORISTA]": limpar_texto_word(motorista), "[CAVALO]": limpar_texto_word(cavalo),
            "[CARRETA]": limpar_texto_word(carreta), "[CONTAINER]": limpar_texto_word(container),
            "[CPF]": limpar_texto_word(cpf), "[DATA]": data_atual
        }
    else:
        nome_arq_base = "VEXA"
        caminho_template = os.path.join(BASE_DIR, "TEMPLATE_COLETA_VEXA.docx")
        tipo_doc = "LOTE"
        numero_doc = lote
        if di: tipo_doc = "DI"; numero_doc = di
        elif bl: tipo_doc = "BL"; numero_doc = bl
            
        substituicoes = {
            "[TIPO_DOC]": tipo_doc, "[NUMERO_DOC]": limpar_texto_word(numero_doc), "[MOTORISTA]": limpar_texto_word(motorista),
            "[CAVALO]": limpar_texto_word(cavalo), "[CARRETA]": limpar_texto_word(carreta),
            "[CONTAINER]": limpar_texto_word(container), "[CPF]": limpar_texto_word(cpf), "[DATA]": data_atual
        }

    if not os.path.exists(caminho_template):
        return None, f"O arquivo '{os.path.basename(caminho_template)}' não foi encontrado."

    try:
        doc = Document(caminho_template)

        for p in doc.paragraphs:
            for key, val in substituicoes.items():
                if key in p.text:
                    for run in p.runs:
                        if key in run.text: run.text = run.text.replace(key, val)
                    if key in p.text: p.text = p.text.replace(key, val)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key, val in substituicoes.items():
                            if key in p.text:
                                for run in p.runs:
                                    if key in run.text: run.text = run.text.replace(key, val)
                                if key in p.text: p.text = p.text.replace(key, val)

        pasta_saida = os.path.join(BASE_DIR, "Ordens_Coleta")
        os.makedirs(pasta_saida, exist_ok=True)
        nome_arquivo_docx = f"Ordem_Coleta_{nome_arq_base}_{limpar_texto_word(container)}_{int(time.time())}.docx"
        nome_arquivo_pdf = f"Ordem_Coleta_{nome_arq_base}_{limpar_texto_word(container)}_{int(time.time())}.pdf"
        caminho_docx = os.path.join(pasta_saida, nome_arquivo_docx)
        caminho_pdf = os.path.join(pasta_saida, nome_arquivo_pdf)
        
        doc.save(caminho_docx)
        time.sleep(1) 
        
        print(Fore.CYAN + f"      🔄 Convertendo a Ordem de Coleta ({nome_arq_base}) para PDF...")
        try:
            convert(caminho_docx, caminho_pdf)
            time.sleep(2) 
            return caminho_pdf, None
        except Exception as e_pdf:
            print(Fore.YELLOW + "      ⚠️ Falha no Word. Limpando memória e tentando segunda vez...")
            os.system("taskkill /F /IM WINWORD.EXE >nul 2>&1")
            time.sleep(3)
            try:
                convert(caminho_docx, caminho_pdf)
                time.sleep(2) 
                return caminho_pdf, None
            except Exception as e2:
                return None, f"Erro ao converter PDF: {e2}"
            
    except Exception as e:
        return None, f"Erro interno Word: {e}"

# ==========================================================
# 🧠 MAPA DE GRUPOS E MOTORISTAS (FROTA)
# ==========================================================
def identificar_grupo_frota(nome_motorista):
    nome = str(nome_motorista).upper().strip()
    if "DJHON" in nome or "DJOHN" in nome: return "001 NORTE NORDESTE"
    if "DAYVSON" in nome: return "002 NORTE NORDESTE"
    if "BRUNO" in nome: return "003 NORTE NORDESTE"
    if "JOSE" in nome or "JOSÉ" in nome or "JERONIMO" in nome or "JERÔNIMO" in nome: return "005 NORTE NORDESTE"
    if "THIAGO" in nome: return "006 NORTE NORDESTE"
    return None

def limpar_tela():
    os.system('cls' if os.name == 'nt' else 'clear')

# ==========================================================
# 📊 LEITURA OMNISCIENTE SELETIVA
# ==========================================================
def obter_aba_atual(service):
    global ABA_CACHEADA, MES_ABA_CACHEADA
    
    # CORREÇÃO BUG 7: Invalida o cache quando o mês muda
    # Antes: ABA_CACHEADA era verificada apenas uma vez e nunca limpa
    mes_agora = datetime.now().month
    if ABA_CACHEADA and MES_ABA_CACHEADA == mes_agora:
        return ABA_CACHEADA
    
    # Mês mudou ou primeira execução: busca a aba correta
    aba_alvo = {
        1: "JANEIRO", 2: "FEVEREIRO", 3: "MARÇO", 4: "ABRIL",
        5: "MAIO", 6: "JUNHO", 7: "JULHO", 8: "AGOSTO",
        9: "SETEMBRO", 10: "OUTUBRO", 11: "NOVEMBRO", 12: "DEZEMBRO"
    }.get(mes_agora, "MARÇO")

    try:
        meta = executar_com_resiliencia_infinita(service.spreadsheets().get(spreadsheetId=PLANILHA_ID))
        abas = [s["properties"]["title"].upper() for s in meta.get("sheets", [])]
        # Aceita MARCO como sinônimo de MARÇO
        if aba_alvo == "MARÇO" and "MARÇO" not in abas and "MARCO" in abas:
            aba_alvo = "MARCO"
        ABA_CACHEADA = aba_alvo if aba_alvo in abas else abas[0]
        MES_ABA_CACHEADA = mes_agora
        return ABA_CACHEADA
    except:
        return aba_alvo

def remover_acentos(texto):
    return ''.join(c for c in unicodedata.normalize('NFD', texto) if unicodedata.category(c) != 'Mn')

def capturar_dados_todas_planilhas_ia(service, aba_alvo, pergunta_usuario):
    pergunta_limpa = remover_acentos(pergunta_usuario.upper())
    portos_solicitados = []
    if "PECEM" in pergunta_limpa: portos_solicitados.append("PECEM")
    if "SANTOS" in pergunta_limpa: portos_solicitados.append("SANTOS")
    if "SALVADOR" in pergunta_limpa: portos_solicitados.append("SALVADOR")
    if "MANAUS" in pergunta_limpa: portos_solicitados.append("MANAUS")
    if "SUAPE" in pergunta_limpa: portos_solicitados.append("SUAPE")
    if any(palavra in pergunta_limpa for palavra in ["TODAS", "TODOS", "GERAL", "TUDO", "COMPLETO", "BRASIL"]):
        portos_solicitados = ["SUAPE", "PECEM", "SANTOS", "SALVADOR", "MANAUS"]
    if not portos_solicitados: portos_solicitados = ["SUAPE"]

    contexto_geral = ""
    for pasta, spread_id in ROTEAMENTO_PORTOS.items():
        nome_porto = os.path.basename(pasta).replace("entrada_", "").upper()
        if nome_porto not in portos_solicitados: continue
            
        print(Fore.CYAN + f"      📥 Fazendo download profundo: {nome_porto} (Aba: {aba_alvo})...")
        contexto_geral += f"\n\n=== DADOS DA PLANILHA DO PORTO: {nome_porto} (MÊS: {aba_alvo}) ===\n"
        try:
            res = executar_com_resiliencia_infinita(service.spreadsheets().get(spreadsheetId=spread_id, ranges=[f"{aba_alvo}!A:AZ"], includeGridData=True))
            grid = res['sheets'][0]['data'][0]
            row_data = grid.get('rowData', [])
            if not row_data: 
                continue
            
            def obter_valor_seguro(dicionario_celula, chave):
                if not isinstance(dicionario_celula, dict): return ""
                val = dicionario_celula.get(chave, "")
                return "" if val is None or str(val).upper() == "NONE" else str(val).strip()
            
            headers = [obter_valor_seguro(cell, 'formattedValue').upper() for cell in row_data[0].get('values', [])]
            linhas_resumo = []
            
            for i in range(1, len(row_data)):
                cells = row_data[i].get('values', [])
                if not cells: continue
                dados_linha = []
                tem_dado_util = False
                for j, cell in enumerate(cells):
                    nome_coluna = headers[j] if j < len(headers) and headers[j] else f"COLUNA_{j+1}"
                    val_str = obter_valor_seguro(cell, 'formattedValue')
                    if val_str: tem_dado_util = True
                    parte = f"{nome_coluna}: {val_str}" if val_str else f"{nome_coluna}: (Vazio)"
                    dados_linha.append(parte)
                if tem_dado_util and dados_linha:
                    linhas_resumo.append(f"📦 REGISTRO {i+1}: " + " | ".join(dados_linha))
            contexto_geral += "\n".join(linhas_resumo)
        except Exception as e:
            print(Fore.RED + f"      ❌ Erro ao baixar {nome_porto}: {e}")
    return contexto_geral

# ==========================================================
# 🧠 MÓDULO DE INTELIGÊNCIA ARTIFICIAL (O CÉREBRO DO ROBÔ)
# ==========================================================
def consultar_gemini(nome_remetente, pergunta_diretoria, contexto_planilhas, caminho_anexo=None):
    data_hoje = datetime.now().strftime("%d/%m/%Y")
    hora_atual = datetime.now().strftime("%H:%M")

    prompt = f"""
    Você é o Jarvis, a Inteligência Artificial Executiva da Norte Nordeste (Empresa de Logística).
    Seu papel é responder perguntas da diretoria de forma educada, proativa e com precisão absoluta.
    
    👤 QUEM ESTÁ PERGUNTANDO: {nome_remetente}
    HOJE É DIA: {data_hoje} e a hora atual é {hora_atual}.
    
    💼 MÓDULO ESPECIAL 1: CRIADOR DE PROPOSTAS COMERCIAIS
    Se o usuário pedir para criar "Proposta" ou "Orçamento":
    COMANDO_PROPOSTA||[CLIENTE]||[CNPJ]||[ORIGEM]||[DESTINO]||[VALOR]||[A/C ou ATT]||[ADVALOREN]||[TRIAGEM]||[PEDAGIO]||[PRAZO DE PAGAMENTO]
    
    💼 MÓDULO ESPECIAL 2: APRESENTAÇÃO DA EMPRESA
    Se o usuário pedir a "apresentação da empresa":
    COMANDO_APRESENTACAO||
    
    💼 MÓDULO ESPECIAL 3: GERADOR DE ORDEM DE COLETA UNIVERSAL
    Se o usuário pedir para gerar "Ordem de Coleta" (ex: para VEXA ou VAL MOTOS), localize o contêiner na BASE DE DADOS EM TEMPO REAL.
    COMANDO_COLETA||[EMPRESA]||[CONTAINER]||[DI]||[LOTE]||[BL]||[MOTORISTA]||[CPF]||[CAVALO]||[CARRETA]
    
    === BASE DE DADOS EM TEMPO REAL ===
    {contexto_planilhas}
    ===================================
    
    Pergunta do {nome_remetente}: "{pergunta_diretoria}"
    
    🛡️ REGRAS DE OURO:
    1. PROATIVIDADE: Sempre informe: *Contêiner, Cliente, Destino e Motorista*.
    2. INVISIBILIDADE: NUNCA escreva a palavra "REGISTRO" na resposta.
    3. TOM DE VOZ: Seja educado e aja como o Jarvis. 
    4. PROIBIDO INVENTAR: Trabalhe APENAS com os dados da BASE DE DADOS.
    5. VISÃO COMPUTACIONAL: Caso eu tenha enviado uma IMAGEM, RECIBO, NOTA FISCAL ou PDF anexo, analise a imagem e responda a pergunta extraindo as informações precisas do documento.
    6. PROIBIDO TEXTO EXTRA NOS COMANDOS: Se a sua resposta for para criar Proposta, Apresentação ou Ordem de Coleta, envie APENAS E EXATAMENTE a linha do comando. Não escreva "Excelente", não dê bom dia. SÓ o comando.
    """

    for i, chave in enumerate(LISTA_CHAVES_GEMINI):
        try:
            client = genai.Client(api_key=chave)
            nome_modelo = obter_modelo_cerebro(client)
            
            print(Fore.MAGENTA + f"      🧠 Processando com o Cérebro Principal ({nome_modelo} - Chave {i+1})...")
            
            contents = [prompt]
            uploaded_file = None
            
            if caminho_anexo and os.path.exists(caminho_anexo):
                print(Fore.MAGENTA + f"      👁️ Enviando mídia (Visão Computacional) para a IA ler...")
                uploaded_file = client.files.upload(file=caminho_anexo)
                contents.append(uploaded_file)
            
            response = client.models.generate_content(
                model=nome_modelo,
                contents=contents,
                config=types.GenerateContentConfig(temperature=0.0) 
            )
            
            if uploaded_file:
                try: client.files.delete(name=uploaded_file.name)
                except: pass
                
            return response.text
        except Exception as e:
            erro_str = str(e).upper()
            if any(termo in erro_str for termo in ["429", "503", "500", "RESOURCE", "UNAVAILABLE", "OVERLOAD"]): 
                print(Fore.YELLOW + f"      ⚠️ Servidor ocupado na Chave {i+1}. Saltando para a próxima...")
            else: 
                print(Fore.YELLOW + f"      ⚠️ Falha na Chave {i+1} ({e}). Saltando para a próxima...")
            continue
            
    return "Desculpe, meus sistemas de processamento estão temporariamente sobrecarregados. Tente em um minuto."

def buscar_entregas_liberadas(service, aba):
    entregas = []
    try:
        def obter_valor_seguro(dicionario_celula, chave):
            if not isinstance(dicionario_celula, dict): return ""
            val = dicionario_celula.get(chave, "")
            return "" if val is None or str(val).upper() == "NONE" else str(val).strip()

        res = executar_com_resiliencia_infinita(
            service.spreadsheets().get(spreadsheetId=PLANILHA_ID, ranges=[f"{aba}!A:AZ"], includeGridData=True)
        )
        grid = res['sheets'][0]['data'][0]
        row_data = grid.get('rowData', [])
        if not row_data: return entregas
        
        headers = [obter_valor_seguro(cell, 'formattedValue').upper() for cell in row_data[0].get('values', [])]
        idx_confirmacao = headers.index("CONFIRMAÇÃO") if "CONFIRMAÇÃO" in headers else 8
        idx_motorista = headers.index("MOTORISTA") if "MOTORISTA" in headers else 6
        idx_destino = headers.index("DESTINO") if "DESTINO" in headers else 9
        idx_container = headers.index("CONTAINER") if "CONTAINER" in headers else 25
                
        for i in range(1, len(row_data)):
            linha_real = i + 1
            cells = row_data[i].get('values', [])
            while len(cells) < max(idx_confirmacao, idx_motorista, idx_destino, idx_container) + 1:
                cells.append({})
            confirmacao = obter_valor_seguro(cells[idx_confirmacao], 'formattedValue').upper()
            if confirmacao in ["OK", "ENVIAR", "ZAP"]:
                motorista = obter_valor_seguro(cells[idx_motorista], 'formattedValue')
                destino_texto = obter_valor_seguro(cells[idx_destino], 'formattedValue')
                destino_nota = obter_valor_seguro(cells[idx_destino], 'note')
                destino_link = obter_valor_seguro(cells[idx_destino], 'hyperlink')
                if destino_link and destino_link not in destino_nota: destino_nota += f" {destino_link}"
                container = obter_valor_seguro(cells[idx_container], 'formattedValue')
                container_limpo = re.sub(r'[^A-Z0-9]', '', container.upper())
                grupo_whatsapp = identificar_grupo_frota(motorista)
                if grupo_whatsapp and container_limpo:
                    entregas.append({
                        "linha": linha_real, "motorista": motorista, "grupo": grupo_whatsapp,
                        "destino": destino_texto, "nota_maps": destino_nota.strip(),
                        "container": container_limpo, "coluna_confirmacao": get_col_letter(idx_confirmacao)
                    })
        return entregas
    except Exception as e:
        print(Fore.RED + f"   ❌ Erro ao buscar entregas: {e}")
        return []

def despachante_continuo():
    limpar_tela()
    print(Fore.BLUE + Style.BRIGHT + "======================================================================")
    print(Fore.BLUE + Style.BRIGHT + "        ROBÔ DESPACHANTE & ASSISTENTE I.A 24/7 - MODO TURBO 2.0")
    print(Fore.BLUE + Style.BRIGHT + "======================================================================\n")
    
    pasta_sessao = os.path.join(BASE_DIR, "WA_Session")
    ultimas_mensagens_ia = {}
    ciclo_frota = 0

    creds = Credentials.from_service_account_file(CREDS_PATH, scopes=SCOPES)
    service = build("sheets", "v4", credentials=creds)

    with sync_playwright() as p:
        browser = p.chromium.launch_persistent_context(
            user_data_dir=pasta_sessao, 
            headless=False,
            viewport={'width': 1280, 'height': 720},
            accept_downloads=True
        )
        page = browser.new_page()
        
        try:
            print(Fore.WHITE + "⏳ A carregar o WhatsApp Web...")
            page.goto("https://web.whatsapp.com/", timeout=120000)
            page.wait_for_selector('div[contenteditable="true"][data-tab="3"]', timeout=300000)
            print(Fore.GREEN + "✅ WhatsApp Conectado e Pronto!\n")
            time.sleep(2)
            
            variacoes_ia = r'IA|I\.A\.?|YA|Y\.A\.?|HIA|ÍA|Í\.A\.?|IH AH|I A|IÁ|I\.Á|EA|E\.A\.?|IY|IYA|IAH|YAH|IHA|JÁ|JA|YÁ|IH|YH|HYA|IIA|IAA|I AH|E AH|E A|IE|I\.E|Y E|YE|YHA|LÁ|LA|EIA|EYA|ILHA|I\.H\.A|Y\.H\.A|I\.Y\.A|I\.E\.A|I\.A\.H|I E A|I E O|I O|I U|Y O'
            variacoes_jarvis = r'JARVIS|JAVIS|JARBES|JARVES|JARVYS|JAVES|JARVI|JABES|JABIS|JAVIZ|JARVIZ|DIARVIS|DJARVIS|TCHARVIS|TJARVIS|YARVIS|YAVIS|JHARVIS|CHARLES|CHAVES|JÁRVIS|XARVIS|XARBES|TCHARLES|GARVIS|GARVES|GARBES|JARIS|JARYS|JARB|JAVI|JAVE|YARVES|YARBES|DJARVES|DIARVES|TCHARVES|CHARVIS|CHARVES|JALVIS|JALVES|JÁVIS|JÁVES|JÁBES|TIARVIS|JHARVES|GABES|GAVIS|MARVIS|MARVES|Jáfes|Javich|Jáfis|JOGOS|JOGO'
            padrao_gatilho = re.compile(rf'^({variacoes_ia}|{variacoes_jarvis})\b[\s,.:;-]*', re.IGNORECASE)

            # Helper para extrair anexos usando clique direito
            def extrair_anexo(msg_node, page_ref):
                try:
                    msg_node.hover()
                    time.sleep(0.5)
                    btn_menu = msg_node.locator('span[data-icon="ic-chevron-down-menu"], span[data-icon="down-context"]').first
                    if btn_menu.is_visible(): btn_menu.click(force=True)
                    else: msg_node.click(button="right", force=True)
                    
                    time.sleep(1)
                    btn_baixar = page_ref.locator('ul li:has-text("Baixar"), ul li:has-text("Download"), div[role="button"]:has-text("Baixar")').first
                    if btn_baixar.is_visible():
                        with page_ref.expect_download(timeout=15000) as download_info:
                            btn_baixar.click(timeout=5000)
                        download = download_info.value
                        ext = os.path.splitext(download.suggested_filename)[1]
                        if not ext: ext = ".png"
                        path = os.path.join(BASE_DIR, f"anexo_in_{int(time.time())}{ext}")
                        download.save_as(path)
                        return path
                    else:
                        page_ref.keyboard.press('Escape')
                except Exception:
                    page_ref.keyboard.press('Escape')
                return None

            while True:
                agora = datetime.now().strftime("%H:%M:%S")
                print(Fore.BLUE + f"\n[{agora}] --- RASTREAMENTO I.A. ---")
                aba_atual = obter_aba_atual(service)
                
                for contato in CONTATOS_DIRETORIA:
                    try:
                        print(Fore.WHITE + f"   🔎 Verificando: {contato}...")
                        busca = page.locator('div[contenteditable="true"][data-tab="3"]')
                        busca.fill("")
                        time.sleep(0.5)
                        busca.fill(contato)
                        time.sleep(1.5)
                        
                        grupo_loc = page.locator(f'span[title="{contato}"]').first
                        if grupo_loc.is_visible():
                            grupo_loc.click()
                            time.sleep(2.0) 
                            
                            todas_mensagens = page.locator('div.message-in, div.message-out').all()
                            
                            if todas_mensagens:
                                ultima_geral = todas_mensagens[-1]
                                is_ultima_nossa = "message-out" in (ultima_geral.get_attribute('class') or "")
                                
                                mensagens_in = [m for m in todas_mensagens if "message-in" in (m.get_attribute('class') or "")]
                                
                                if mensagens_in:
                                    ultima_in = mensagens_in[-1]
                                    msg_id = ultima_in.get_attribute("data-id")
                                    if not msg_id: msg_id = ultima_in.inner_text().strip()
                                        
                                    msg_antiga = ultimas_mensagens_ia.get(contato, "")
                                    
                                    if msg_id != msg_antiga:
                                        ultimas_mensagens_ia[contato] = msg_id
                                        
                                        # Apenas recusa continuar se a ÚLTIMA mensagem for do robô
                                        if is_ultima_nossa:
                                            print(Fore.WHITE + f"   ⏭️ Sincronizado: O último pedido de {contato} já foi respondido.")
                                            continue
                                        
                                        texto_puro_tela = ultima_in.inner_text().strip()
                                        linhas_msg = texto_puro_tela.split('\n')
                                        if linhas_msg and re.search(r'\d{2}:\d{2}', linhas_msg[-1]): 
                                            texto_msg_limpo = "\n".join(linhas_msg[:-1]).strip()
                                        else:
                                            texto_msg_limpo = "\n".join(linhas_msg).strip()

                                        texto_para_processar = texto_msg_limpo
                                        caminho_anexo = None
                                        is_audio = False
                                        
                                        # Verifica se a última mensagem é mídia
                                        is_audio_html = ultima_in.locator('span[data-icon="audio-play"], span[data-icon="ptt-status"], span[data-icon="audio-download"], audio').count() > 0
                                        is_audio_regex = bool(re.search(r'\d{1,2}:\d{2}\s*\d[.,]\d[xX×]?', texto_puro_tela))
                                        
                                        if is_audio_html or is_audio_regex:
                                            is_audio = True
                                            print(Fore.CYAN + f"   🎤 Áudio detectado! Iniciando interceptação...")
                                            try:
                                                btn_load = ultima_in.locator('span[data-icon="audio-download"]')
                                                if btn_load.count() > 0 and btn_load.first.is_visible():
                                                    btn_load.first.click()
                                                    time.sleep(2.5)
                                                    
                                                caminho_audio = extrair_anexo(ultima_in, page)
                                                if caminho_audio:
                                                    transcricao = transcrever_audio_local(caminho_audio)
                                                    if transcricao: 
                                                        texto_para_processar = transcricao
                                                        print(Fore.YELLOW + f"   🗣️ Transcrição: '{transcricao}'")
                                            except Exception as e:
                                                print(Fore.RED + f"   ❌ Erro áudio: {e}")
                                        else:
                                            tem_midia = ultima_in.locator('img, span[data-icon="document"]').count() > 0
                                            if tem_midia:
                                                print(Fore.CYAN + "   📸 Imagem/Documento detectado na última mensagem!")
                                                caminho_anexo = extrair_anexo(ultima_in, page)
                                                
                                        # Se não achou mídia na última, olha APENAS a mídia da anterior (NÃO o texto)
                                        if not is_audio and not caminho_anexo and len(mensagens_in) >= 2:
                                            msg_anterior = mensagens_in[-2]
                                            tem_midia_ant = msg_anterior.locator('img, span[data-icon="document"]').count() > 0
                                            if tem_midia_ant:
                                                print(Fore.CYAN + "   📸 Imagem/Documento detectado na mensagem anterior!")
                                                caminho_anexo = extrair_anexo(msg_anterior, page)

                                        # =======================================================
                                        # 🔥 FLUXO DE EXECUÇÃO
                                        # =======================================================
                                        if texto_para_processar:
                                            gatilho_ativado = False
                                            pergunta_limpa = ""
                                            linhas_originais = texto_para_processar.split('\n')
                                            
                                            for i in range(len(linhas_originais)-1, -1, -1):
                                                linha_atual = linhas_originais[i].strip()
                                                if padrao_gatilho.search(linha_atual):
                                                    gatilho_ativado = True
                                                    bloco_texto = "\n".join(linhas_originais[i:])
                                                    pergunta_limpa = padrao_gatilho.sub('', bloco_texto).strip()
                                                    break
                                            
                                            if gatilho_ativado:
                                                nome_pessoa = contato.split('-')[0].replace('NN', '').strip()
                                                if "Financeiro" in contato: nome_pessoa = "Diego"
                                                
                                                print(Fore.YELLOW + f"   🤖 PEDIDO RECEBIDO: '{pergunta_limpa}'")
                                                
                                                caixa_texto_envio = page.locator('footer div[contenteditable="true"]').first
                                                
                                                if any(palavra in pergunta_limpa.upper() for palavra in ["APRESENTAÇÃO", "APRESENTACAO", "PORTFÓLIO", "PORTFOLIO"]):
                                                    dados_planilha = "Ignorar planilha, enviar apresentação."
                                                elif any(palavra in pergunta_limpa.upper() for palavra in ["PROPOSTA", "ORÇAMENTO", "ORCAMENTO"]):
                                                    dados_planilha = "Ignorar planilha, gerar proposta."
                                                else:
                                                    aba_pesquisa = aba_atual
                                                    meses_lista = ["JANEIRO", "FEVEREIRO", "MARÇO", "MARCO", "ABRIL", "MAIO", "JUNHO", "JULHO", "AGOSTO", "SETEMBRO", "OUTUBRO", "NOVEMBRO", "DEZEMBRO"]
                                                    for mes in meses_lista:
                                                        if mes in pergunta_limpa.upper():
                                                            aba_pesquisa = "MARÇO" if mes == "MARCO" else mes
                                                            break
                                                    dados_planilha = capturar_dados_todas_planilhas_ia(service, aba_pesquisa, pergunta_limpa)
                                                
                                                resposta_ia = consultar_gemini(nome_pessoa, pergunta_limpa, dados_planilha, caminho_anexo)
                                                
                                                if "COMANDO_APRESENTACAO||" in resposta_ia:
                                                    caminho_apres = os.path.join(BASE_DIR, "APRESENTACAO_COMERCIAL.pdf")
                                                    if not os.path.exists(caminho_apres):
                                                        caixa_texto_envio.fill(f"🤖 *Jarvis*\nDesculpe {nome_pessoa}, não encontrei o ficheiro 'APRESENTACAO_COMERCIAL.pdf'.")
                                                        page.keyboard.press('Enter')
                                                    else:
                                                        try:
                                                            copiar_arquivos_windows(caminho_apres)
                                                            caixa_texto_envio.click()
                                                            time.sleep(1)
                                                            page.keyboard.press('Control+V')
                                                            time.sleep(3)
                                                            page.locator('span[data-icon="send"], span[data-icon="wds-ic-send-filled"]').first.click(timeout=10000)
                                                            time.sleep(3)
                                                            caixa_texto_envio.fill(f"🤖 *Jarvis*\n\n✅ {nome_pessoa}, aqui está a nossa Apresentação Comercial Oficial!")
                                                            page.keyboard.press('Enter')
                                                        except Exception as e:
                                                            # CORREÇÃO BUG 16: loga o erro em vez de silenciar
                                                            print(Fore.RED + f"      ❌ Erro ao enviar apresentação: {e}")

                                                elif "COMANDO_PROPOSTA||" in resposta_ia:
                                                    linha_comando = next(linha for linha in resposta_ia.split('\n') if "COMANDO_PROPOSTA||" in linha)
                                                    partes = linha_comando.split("||")
                                                    
                                                    cliente_p = partes[1].strip() if len(partes) > 1 else "ERRO"
                                                    cnpj_p = partes[2].strip() if len(partes) > 2 else ""
                                                    origem_p = partes[3].strip() if len(partes) > 3 else "A DEFINIR"
                                                    destino_p = partes[4].strip() if len(partes) > 4 else "A DEFINIR"
                                                    valor_p = partes[5].strip() if len(partes) > 5 else "A DEFINIR"
                                                    att_p = partes[6].strip() if len(partes) > 6 else "A DEFINIR"
                                                    advaloren_p = partes[7].strip() if len(partes) > 7 else "A DEFINIR"
                                                    triagem_p = partes[8].strip() if len(partes) > 8 else "A DEFINIR"
                                                    pedagio_p = partes[9].strip() if len(partes) > 9 else "A DEFINIR"
                                                    prazo_p = partes[10].strip() if len(partes) > 10 else "A DEFINIR"
                                                    
                                                    if cliente_p != "ERRO":
                                                        caminho_doc_pdf, erro = gerar_documento_proposta(cliente_p, cnpj_p, origem_p, destino_p, valor_p, att_p, advaloren_p, triagem_p, pedagio_p, prazo_p)
                                                        if erro:
                                                            caixa_texto_envio.fill(f"🤖 *Jarvis*\nErro: {erro}")
                                                            page.keyboard.press('Enter')
                                                        else:
                                                            try:
                                                                copiar_arquivos_windows(caminho_doc_pdf)
                                                                caixa_texto_envio.click()
                                                                time.sleep(1)
                                                                page.keyboard.press('Control+V')
                                                                time.sleep(3) 
                                                                page.locator('span[data-icon="send"], span[data-icon="wds-ic-send-filled"]').first.click()
                                                                time.sleep(3)
                                                                caixa_texto_envio.fill(f"🤖 *Jarvis*\n\n✅ {nome_pessoa}, proposta gerada!")
                                                                page.keyboard.press('Enter')
                                                            except Exception as e:
                                                                # CORREÇÃO BUG 16: loga o erro em vez de silenciar
                                                                print(Fore.RED + f"      ❌ Erro ao enviar proposta: {e}")

                                                elif "COMANDO_COLETA||" in resposta_ia:
                                                    linha_comando = next(linha for linha in resposta_ia.split('\n') if "COMANDO_COLETA||" in linha)
                                                    partes = linha_comando.split("||")
                                                    
                                                    empresa_c = partes[1].strip() if len(partes) > 1 else "VEXA"
                                                    container_c = partes[2].strip() if len(partes) > 2 else "ERRO"
                                                    di_c = partes[3].strip() if len(partes) > 3 else ""
                                                    lote_c = partes[4].strip() if len(partes) > 4 else ""
                                                    bl_c = partes[5].strip() if len(partes) > 5 else ""
                                                    motorista_c = partes[6].strip() if len(partes) > 6 else ""
                                                    cpf_c = partes[7].strip() if len(partes) > 7 else ""
                                                    cavalo_c = partes[8].strip() if len(partes) > 8 else ""
                                                    carreta_c = partes[9].strip() if len(partes) > 9 else ""
                                                    
                                                    if container_c != "ERRO":
                                                        caminho_coleta_pdf, erro_coleta = gerar_ordem_coleta(empresa_c, container_c, di_c, lote_c, bl_c, motorista_c, cpf_c, cavalo_c, carreta_c)
                                                        if erro_coleta:
                                                            caixa_texto_envio.fill(f"🤖 *Jarvis*\nDesculpe {nome_pessoa}, ocorreu um erro interno: {erro_coleta}")
                                                            page.keyboard.press('Enter')
                                                        else:
                                                            try:
                                                                copiar_arquivos_windows(caminho_coleta_pdf)
                                                                caixa_texto_envio.click()
                                                                time.sleep(1)
                                                                page.keyboard.press('Control+V')
                                                                time.sleep(3) 
                                                                page.locator('span[data-icon="send"], span[data-icon="wds-ic-send-filled"]').first.click(timeout=10000)
                                                                time.sleep(3)
                                                                
                                                                texto_final = f"🤖 *Jarvis (I.A. Norte Nordeste)*\n\n✅ {nome_pessoa}, a Ordem de Coleta do contêiner *{container_c}* ({empresa_c}) foi gerada com sucesso!\n\n🚚 Motorista: {motorista_c}\nPlacas: {cavalo_c} / {carreta_c}"
                                                                caixa_texto_envio.fill(texto_final)
                                                                page.keyboard.press('Enter')
                                                            except Exception as e:
                                                                # CORREÇÃO BUG 16: loga o erro em vez de silenciar
                                                                print(Fore.RED + f"      ❌ Erro ao enviar ordem de coleta: {e}")
                                                    else:
                                                        caixa_texto_envio.fill(f"🤖 *Jarvis*\n{nome_pessoa}, não consegui entender os dados do contêiner.")
                                                        page.keyboard.press('Enter')

                                                else:
                                                    texto_final = f"🤖 *Jarvis (I.A. Norte Nordeste)*\n\n{resposta_ia}"
                                                    caixa_texto_envio.click()
                                                    page.keyboard.insert_text(texto_final)
                                                    time.sleep(0.5)
                                                    page.keyboard.press('Enter')
                                                    time.sleep(2)
                                                    
                                                if caminho_anexo and os.path.exists(caminho_anexo):
                                                    try: os.remove(caminho_anexo)
                                                    except: pass
                                            else:
                                                print(Fore.WHITE + f"   ⏭️ Ignorado (Não possui o gatilho na mensagem).")
                            else:
                                print(Fore.WHITE + f"   👁️ Chat vazio.")
                    except Exception as e:
                        print(Fore.RED + f"   ⚠️ Erro ao verificar {contato}: {e}")

                # ======================================================
                # 🚚 FASE 2: DESPACHO DE DOCUMENTAÇÃO (A CADA ~1 MINUTO)
                # ======================================================
                ciclo_frota += 1
                if ciclo_frota >= 4:
                    print(Fore.WHITE + f"\n🚚 [{agora}] Verificando despachos da frota...")
                    entregas = buscar_entregas_liberadas(service, aba_atual)
                    
                    if entregas:
                        print(Fore.GREEN + f"🔔 {len(entregas)} nova(s) entrega(s) liberada(s)!")
                        for entrega in entregas:
                            cont = entrega["container"]
                            grupo = entrega["grupo"]
                            
                            pasta_cont = encontrar_pasta_container(cont)
                            pdfs = [os.path.join(pasta_cont, f) for f in os.listdir(pasta_cont) if f.lower().endswith('.pdf')] if pasta_cont and os.path.exists(pasta_cont) else []
                            
                            if not pdfs:
                                try: executar_com_resiliencia_infinita(service.spreadsheets().values().update(spreadsheetId=PLANILHA_ID, range=f"{aba_atual}!{entrega['coluna_confirmacao']}{entrega['linha']}", valueInputOption="USER_ENTERED", body={"values": [["FALTA PDF"]]}))
                                except: pass
                                continue
                                
                            busca = page.locator('div[contenteditable="true"][data-tab="3"]')
                            busca.fill(grupo)
                            time.sleep(2)
                            grupo_loc = page.locator(f'span[title="{grupo}"]').first
                            if not grupo_loc.is_visible(): continue
                            grupo_loc.click()
                            time.sleep(2)
                            
                            texto_msg = f"🚚 *NOVA ENTREGA LIBERADA*\n\n📦 *Contêiner:* {cont}\n📍 *Destino:* {entrega['destino']}"
                            if entrega['nota_maps']: texto_msg += f"\n🗺️ *Localização/Google Maps:* {entrega['nota_maps']}"
                            
                            try:
                                copiar_arquivos_windows(pdfs)
                                caixa_texto_envio = page.locator('footer div[contenteditable="true"]').first
                                caixa_texto_envio.click()
                                time.sleep(1)
                                page.keyboard.press('Control+V')
                                time.sleep(3)
                                page.locator('span[data-icon="send"], span[data-icon="wds-ic-send-filled"]').first.click(timeout=10000)
                                time.sleep(4) 
                                
                                caixa_texto_envio.fill(texto_msg)
                                page.keyboard.press('Enter')
                                
                                executar_com_resiliencia_infinita(service.spreadsheets().values().update(spreadsheetId=PLANILHA_ID, range=f"{aba_atual}!{entrega['coluna_confirmacao']}{entrega['linha']}", valueInputOption="USER_ENTERED", body={"values": [["ENVIADO"]]}))
                            except Exception as e:
                                print(Fore.RED + f"      ❌ Erro frota: {e}")
                    ciclo_frota = 0
                            
                print(Fore.CYAN + f"⏳ Retornando ao modo escuta (5s)...")
                time.sleep(5)
                
        except Exception as e:
            print(Fore.RED + f"\n❌ Erro Crítico no WhatsApp: {e}")
        finally:
            browser.close()

if __name__ == "__main__":
    despachante_continuo()